from flask import Flask, request, render_template
from docx import Document
import spacy
import os
import re
from googlesearch import search

app = Flask(__name__)
nlp = spacy.load("en_core_web_sm")
UPLOAD_FOLDER = "uploads"
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
web_search_allowed = False

def read_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []
        youtube_links = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                links = re.findall(r'(https?://(?:www\.)?youtube\.com/\S+)', para.text)
                youtube_links.extend(links)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
                        links = re.findall(r'(https?://(?:www\.)?youtube\.com/\S+)', cell.text)
                        youtube_links.extend(links)
        return "\n".join(full_text), youtube_links
    except Exception as e:
        return f"Error reading document: {str(e)}", []

def answer_question(doc_text, question, youtube_links):
    global web_search_allowed
    if not doc_text or not question:
        return "Arre bhai, file ya question toh daal! Kya main hawa se jawab du? 😜", []
    doc = nlp(doc_text)
    question_doc = nlp(question)
    best_match = None
    highest_similarity = 0.0
    relevant_links = []
    for sent in doc.sents:
        similarity = sent.similarity(question_doc)
        if similarity > highest_similarity:
            highest_similarity = similarity
            best_match = sent.text
        for link in youtube_links:
            if link in sent.text and link not in relevant_links:
                relevant_links.append(link)
    if best_match and highest_similarity > 0.3:
        response = f"Haan bhai, sun na! Tera jawab yeh hai: {best_match} 😎"
        if relevant_links:
            response += "\nAur yeh lo YouTube links for reference:\n" + "\n".join([f"- {link}" for link in relevant_links])
        product_example = get_indian_product(question)
        if product_example:
            response += f"\n\nIndian example, boss: {product_example}"
        return response, relevant_links
    response = "Arre, yeh toh image ya equation jaisa lagta hai! Notes check kar, ya main web surf karu? Thodi si unreliable info ho sakti hai. 😅 Bolo, 'yes' toh main search karta hu!"
    web_search_allowed = True
    return response, []

def get_indian_product(question):
    question_lower = question.lower()
    if "engine" in question_lower or "vehicle" in question_lower:
        return "Tata Nexon EV - ekdum zabardast electric SUV, Indian style! 🇮🇳"
    elif "refrigerator" in question_lower or "air conditioner" in question_lower:
        return "Godrej ka fridge ya Voltas ka AC - garmi bhagao, Indian swag! 😎"
    elif "robotics" in question_lower or "automation" in question_lower:
        return "GreyOrange ke warehouse robots - Indian tech ka dhamaka! 🤖"
    else:
        try:
            query = f"Indian products related to {question}"
            for url in search(query, num_results=1):
                return f"Check products like on {url} - Indian market ka scene! 🌟"
        except:
            return None
    return None

@app.route("/", methods=["GET", "POST"])
def index():
    global web_search_allowed
    response = "Comps ka banda Mech bol raha hun! 😎 .docx file daal aur question pooch, main mast Hinglish mein jawab deta hu!"
    youtube_links = []
    if request.method == "POST":
        if "file" not in request.files:
            response = "Bhai, file toh upload kar! Main kya khali haath jawab du? 😜"
        else:
            file = request.files["file"]
            if file.filename == "":
                response = "Arre, koi file select kar pehle! 😅"
            elif file and file.filename.endswith(".docx"):
                file_path = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
                file.save(file_path)
                doc_text, youtube_links = read_docx(file_path)
                question = request.form.get("question", "")
                if question.lower() == "yes" and web_search_allowed:
                    try:
                        web_results = []
                        for url in search(question, num_results=3):
                            web_results.append(url)
                        response = f"Web se yeh lo, bhai, thodi unreliable ho sakti hai: \n" + "\n".join([f"- {url}" for url in web_results])
                        web_search_allowed = False
                    except:
                        response = "Oops, web search mein kuch gadbad ho gaya! Notes hi check kar le. 😅"
                elif question:
                    response, youtube_links = answer_question(doc_text, question, youtube_links)
                else:
                    response = "File toh upload ho gaya, ab question bhi pooch le! 😎"
            else:
                response = "Sirf .docx file chalega, bhai! Yeh kya daal diya? 😜"
    return render_template("index.html", response=response, youtube_links=youtube_links)

if __name__ == "__main__":
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    app.run(debug=True)